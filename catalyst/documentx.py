# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches
import tempfile
import os
import urllib2
import MultipartPostHandler
import flask
import os
import sys
import fdfgen
import copy


def fix_qty(function):
    def wrapper(zohoString, *args, **kwargs):
        tabledicts = function(zohoString)
        for tabledict in tabledicts:
            tabledict["qty"] = tabledict["qty"].split(".")[0]
        return tabledicts
    return wrapper




def datefix(date):
    try:
        datefield = date.split("-")
        datefixed = "/".join([datefield[1], datefield[2], datefield[0]])
    except:
        datefixed = ""
    return datefixed


def calculate_total(ptable):
    allow_sum, retail_sum = 0, 0
    for rows in ptable:
        qty, allowable, retail = float(rows["qty"]), float(
            rows["allowable"]), float(rows["price"])
        allow_sum += qty * allowable
        retail_sum += qty * retail
    return retail_sum, allow_sum


def upload_with_temp(doc, ID, pfix="trumobilitytest"):
    temp_doc, filename = tempfile.mkstemp(
        prefix=pfix, suffix=".docx")
    try:
        doc.save(filename)
        os.close(temp_doc)
        zoho_upload(filename, ID)
    finally:
        os.remove(filename)


@fix_qty
def createMapFromZohoString(zohoString):
    if createMapFromZohoString == "":
        return None
    ItemList = zohoString.split("-")
    zss = zohoString.split("<:>")
    sanlist = []
    sanlist2 = []
    for items in zss:
        sanlist += items.split("\t")
    for items in sanlist:
        sanlist2 += items.split("\n")
    
    sanlist2 = [san for san in sanlist2 if san!=u'']
    rlist = []
    fdict = {}
    for i in range(0,len(sanlist2),2):
        if sanlist2[i] == u'partno':
            fdict[sanlist2[i]] = sanlist2[i+1]
            rlist.append(copy.deepcopy(fdict))
            fdict = {}
            continue
        fdict[sanlist2[i]] = sanlist2[i+1]
    for entry in rlist:
        entry['total'] = '{0:.2f}'.format(float(entry['price']) * float(entry['qty']))
        entry['totalallow'] = '{0:.2f}'.format(float(entry['allowable']) * float(entry['qty']))
        print entry
    
    return rlist


def zoho_upload(doc, ID, authtoken="beddaafb803d5a9cd868ad43ae6756ff"):
    opener = urllib2.build_opener(MultipartPostHandler.MultipartPostHandler)
    params = {'authtoken': authtoken, 'scope': 'crmapi', 'newFormat': '1',
              'id': str(ID).encode("utf-8"), 'content': open(doc, "rb")}
    finalurl = "https://crm.zoho.com/crm/private/xml/Potentials/uploadFile"
    results = opener.open(finalurl, params)
    print results.read()
    return results.read()


def build_form_from_fields(form, fields, rl, headers):
    document = Document(form)

    for label in document.paragraphs:
        # print label.text.strip()
        if label.text.strip() in fields.keys():
            label.add_run(":   {}".format(fields[label.text.strip()]))

    try:
        table = document.tables[0]
        while len(table.rows) < len(rl):
            table.add_row()
        for out_index, row in enumerate(table.rows):
            for index, head in enumerate(headers):
                try:
                    row.cells[index].text = rl[out_index].get(head, "")
                except:
                    pass
    except:
        pass
    return document


def DeliveryTicket(rform):
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    total_tup = calculate_total(rl)
    fields = {"Phone": rform.get('patientphone', ""),
              "Patient Name": rform.get('patientname', ""),
              "Address": rform.get('patientaddress', ""),
              "DOB": datefix(rform['patientdob']),
              "Totals": "{}  {}".format(*total_tup)}
    print fields
    headers = ["qty", "hcpcs", "name", "manufacturer",
               "model", "partno", "total", "totalallow"]
    delivery_ticket = build_form_from_fields(
        "delivery_ticket.docx", fields, rl, headers)
    upload_with_temp(delivery_ticket, rform.get(
        "poid", "2282460000002245053"), pfix="Delivery_Ticket_")


def DetailedProductDescription(rform):
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    fields = {"Phone": rform.get('patientphone', ""),
              "Patient Name": rform.get('patientname', ""),
              "Address": rform.get('patientaddress', ""),
              "DOB": datefix(rform['patientdob']),
              "Physician Name": rform.get('physicianname', ""),
              "Physician Address": rform.get('physicianaddress', ""),
              "Physician Phone": rform.get('physicianphone', ""),
              "Fax": rform.get('physicianfax', ""),
              "NPI number": rform.get('physiciannpi', "")}
    headers = ["qty", "hcpcs", "name", "manufacturer", "model", "partno"]
    detailedpd = build_form_from_fields(
        "Detailed_Product_Description.docx", fields, rl, headers)
    upload_with_temp(detailedpd, rform.get("poid"),
                     pfix="Detailed_Product_Description_")


def EquipmentMeasurementSheet(rform):
    # rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    fields = {"Primary Ins :": rform.get('patientprimaryinsurance', ""),
              "Secondary Ins :": rform.get('patientsecondaryinsurance', ""),
              "Dx codes :": rform.get('dxcode', ""),
              "ATP :": rform.get('atpname', ""),
              "Facility :": rform.get('facilityname', ""),
              "Name :": rform.get('patientname', ""),
              "DOB :": datefix(rform['patientdob']),
              "Diagnosis :": rform.get('patientallicds', "")}
    headers = []
    # detailedpd = build_form_from_fields(
    #     "Equipment_Measurement_Sheet.docx", fields, rl, headers)
    document = Document("Equipment_Measurement_Sheet.docx")
    # style = document.styles['Normal']
    # font = style.font
    # font.underline = True
    for label in document.paragraphs:
        # print label.text.strip()
        if label.text.strip() in fields.keys():
            # print "hi"
            print label.text.strip()
            # label.add_run(":  ")
            label.add_run("{}".format(fields[label.text.strip()])).underline = True

    upload_with_temp(document, rform.get("poid"),
                     pfix="Equipment_Measurement_Sheet_")


def WheelchairEval(rform):
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    fields = {u"CLIENT’S NAME": rform.get('patientname', ""),
              u"CLIENT’S ADDRESS": rform.get('patientaddress', ""),
              u"CLIENT’S DATE OF BIRTH": datefix(rform['patientdob']),
              u"ICD-10 CODE(S)": rform.get('patienticds', ""),
              u"CLIENT’S MEDICAID #": rform.get('patientmedicaid', ""),
              "HEIGHT": rform.get('patientheight', ""),
              "Therapist Name": rform.get('therapistname', ""),
              "Phone": rform.get('therapistphone', ""),
              "Title": rform.get('therapisttitle', ""),
              "Employer Name": rform.get('therapistemp', ""),
              "WEIGHT": rform.get('patientweight', "")}
    headers = ["qty", "hcpcs", "name", "justification"]
    detailedpd = build_form_from_fields(
        "UT_MCD_wheelchair_eval_fom.docx", fields, rl, headers)
    upload_with_temp(detailedpd, rform.get("poid"), pfix="Wheelchair_Eval_")


def MedicaidDelivery(rform):
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    fields = {"Patient Name": rform.get('patientname', ""),
              "Patient / Responsible Party Name": rform.get('patientname', "")}
    headers = []
    detailedpd = build_form_from_fields(
        "Medicaid_Final_Evaluation.docx", fields, rl, headers)
    upload_with_temp(detailedpd, rform.get("poid"),
                     pfix="Medicaid_Final_Evaluation_")


def PowerChairTrainingChecklist(rform):
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    fields = {"PATIENT NAME": (40 * " ") + rform.get('patientname', "") + (40 * " ") + "DATE:",
              "THERAPIST WHO OBSERVED THE TRAINING": rform.get('therapistname', "")}
    headers = []
    detailedpd = build_form_from_fields(
        "POWER_CHAIR_TRAINING_CHECKLIST.docx", fields, rl, headers)
    upload_with_temp(detailedpd, rform.get("poid"),
                     pfix="POWER_CHAIR_TRAINING_CHECKLIST_")


def PriorAuth(rform):
    all_fields = FirstPageFDF(rform)
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("file_fdf.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk utahPA.pdf fill_form file_fdf.fdf output Medicaid_PA.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    if len(rl) > 3:
        index = 3
        while index < len(rl):
            OverflowFDF(rl, start=index)
            index += 18
    zoho_upload("Medicaid_PA.pdf", rform.get("poid"))

def writeFDFoverflow(all_fields):
    print "ALL FIELDS:\n"
    print all_fields
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("overflowfile_fdf.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk utahPA_overflow.pdf fill_form overflowfile_fdf.fdf output Medicaid_overflow.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)



def FirstPageFDF(rform):
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    headers = ["Code Description {}", "16 CPTHCPCS{}", "17 MODIFIER{}", "18 UNITSVISITS{}", "19 ESTIMATED COST{}"]
    zheaders = ["name", "hcpcs", "modifier", "qty", "price"]
    all_fields_key = [("2 MEDICAID MEMBER NAME", 'patientname'), ("Medicaid Member ID", 'patientmedicaid'),
    ("Medicaid Member ID", 'patientmedicaid'), ("20 DIAGNOSIS DESCRIPTION  ICD10CM CODES 2", 'patienticds'),
    ( "Hospital Address", 'facilitystreet'), ("NPI_2", 'facilitynpi'), ("Facility Name", 'facilityname'),
    ("Phone", 'facilityphone'),("NAME_3", 'physicianname'), ("Prescriber phone", 'physicianphone'),
    ("NAME_2", 'facilityname'),("Hospital phone", 'facilityphone'),("Hospital fax", 'facilityfax')]
    all_fields = [(x[0], rform.get(x[1], "")) for x in all_fields_key]
    all_fields += [("4 AGE", rform.get('patientage').split(".")[0])]
    all_fields += [("3 DATE OF BIRTH", datefix(rform.get('patientdob')))]
    all_fields += [("Date of Request", datefix(rform.get('today')))]
    firstpage_items = rl[:3]
    #for items in firstpage_items:
    #    print items
    for index, record in enumerate(firstpage_items):
        #print index
        #print record
        fill_headers = map(lambda x: x.format(index + 1), headers)
        total = '{0:.2f}'.format(float(record.get("qty")) * float(record.get("price")))
        fill_kv = zip(fill_headers, [record.get(head, "") for head in zheaders if head!="price"])
        fill_kv.append(("19 ESTIMATED COST{}".format(index + 1), total))
        all_fields += fill_kv
    return all_fields

def OverflowFDF(rl, start=4):
    all_fields = []
    headers = ["Code Description{}", "16 CPTHCPCS{}", "17 MODIFIER{}", "18 UNITSVISITS{}", "19 ESTIMATED COST{}"]
    zheaders = ["name", "hcpcs", "modifier", "qty", "price"]
    indexes = range(start, start+18)
    fill_kv = []
    for idx, index in enumerate(indexes):
        try:
            record = rl[index]
            total = '{0:.2f}'.format(float(record.get("qty")) * float(record.get("price")))
            fill_headers = map(lambda x: x.format(idx + 4), headers)
            fill = zip(fill_headers, [rl[index].get(head, "") for head in zheaders  if head!="price"])
            fill.append(("19 ESTIMATED COST{}".format(idx + 4), total))
            fill_kv += fill
        except:
            break
    writeFDFoverflow(fill_kv)
    mergecom = "/app/vendor/pdftk/bin/pdftk Medicaid_PA.pdf Medicaid_overflow.pdf cat output Medicaid_PA_final.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(mergecom)
    os.rename("Medicaid_PA_final.pdf", "Medicaid_PA.pdf")

def idaho_mcd_pa(rform):

    # Map fields
    all_fields_key = [("last_name", 'patient_last_name'), ("first_name", "patient_first_name"), ("medicaid_id", "patientmedicaid"),
    ("provider_name", "patientprimaryinsurance"), ("npi", "physiciannpi"), 
    ("contact_person", "primarycontact"), ("phone", "providerphone")]

    # Fill Data
    all_fields = [(x[0], rform.get(x[1], "")) for x in all_fields_key]
    all_fields += [("dob", datefix(rform.get('patientdob')))]

    # Fill Requested Equipment
    headers = ["hcpcs{}", "description{}", "qty{}", "price{}"]
    zheaders = ["hcpcs", "name", "qty", "price"]
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    firstpage_items = rl[:9]
    for index, record in enumerate(firstpage_items):
        fill_headers = map(lambda x: x.format(index + 1), headers)
        fill_kv = zip(fill_headers, [record.get(head, "") for head in zheaders])
        all_fields += fill_kv

    # Convert
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("forms/idaho_mcd_pa_data.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk forms/idaho_mcd_pa.pdf fill_form forms/idaho_mcd_pa_data.fdf output Idaho_MCD_DME_Wheelchair_Repair_PA.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)
    zoho_upload("Idaho_MCD_DME_Wheelchair_Repair_PA.pdf", rform.get("poid"))

# NV Wheel chair repair form
def nv_wheel_repair(rform):

    # Map fields
    all_fields_key = [("patient_name", 'patientnamelf'), ("pat_mcd_id", "patientmedicaid"), ("pat_phone", "patientphone"),
    ("provider_name", "patientprimaryinsurance"), ("npi", "physiciannpi"),("provider_phone", "providerphone")]

    # Fill Data
    all_fields = [(x[0], rform.get(x[1], "")) for x in all_fields_key]
    all_fields += [("dob", datefix(rform.get('patientdob')))]

    # Convert
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("forms/nv_wheelchair_repair_data.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk forms/nv_wheelchair_repair.pdf fill_form forms/nv_wheelchair_repair_data.fdf output NV_Wheel_Chair_Repair_Form.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)
    zoho_upload("NV_Wheel_Chair_Repair_Form.pdf", rform.get("poid"))

# ID DMERequest Form
def ID_DMERequest(rform):

    # Map fields
    all_fields_key = [("last_name", 'patient_last_name'), ("first_name", "patient_first_name"), ("medicaid_id", "patientmedicaid"),
    ("provider_name", "patientprimaryinsurance"), ("npi", "physiciannpi"), ("contact_person", "primarycontact"), ("provider_email", "providerphone"),
    ("phy_name", "physicianname"), ("phy_phone", "physicianphone")]

    # Fill Data
    all_fields = [(x[0], rform.get(x[1], "")) for x in all_fields_key]
    all_fields += [("dob", datefix(rform.get('patientdob')))]

    # Fill Requested Equipment
    headers = ["hcpcs{}", "description{}", "qty{}", "price{}"]
    zheaders = ["hcpcs", "name", "qty", "price"]
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    firstpage_items = rl[:6]
    for index, record in enumerate(firstpage_items):
        fill_headers = map(lambda x: x.format(index + 1), headers)
        fill_kv = zip(fill_headers, [record.get(head, "") for head in zheaders])
        all_fields += fill_kv

    # Convert
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("forms/ID_DMERequestForm_data.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk forms/ID_DMERequestForm.pdf fill_form forms/ID_DMERequestForm_data.fdf output ID_DMERequestForm.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)
    zoho_upload("ID_DMERequestForm.pdf", rform.get("poid"))


# Mobility Evaluation Template Idaho
def mobility_evaluation_template_idaho(rform):

    # Map fields
    all_fields_key = [("patient_name", 'patientnamelf'), ("MID", "patientmedicaid"), ("Address", "patientaddress"),
    ("Age", "patientage"), ("Phone", "patientphone"), ("Physician", "physicianname"), ("Weight", "patientweight"), ("Height", "patientheight"),
    ("ICD", "patienticd"), ("ICD_2", "patienticd2"), ("ICD_3", "patienticd3")]

    # Fill Data
    all_fields = [(x[0], rform.get(x[1], "")) for x in all_fields_key]

    # Convert
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("forms/MobilityEvaluationTemplate_Idaho_data.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk forms/MobilityEvaluationTemplate_Idaho.pdf fill_form forms/MobilityEvaluationTemplate_Idaho_data.fdf output MobilityEvaluationTemplate_Idaho.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)
    zoho_upload("MobilityEvaluationTemplate_Idaho.pdf", rform.get("poid"))

# NV DME Prior auth form
def nv_dme_prior_auth(rform):

    # Map fields
    all_fields_key = [("patient_name", 'patientnamelf'), ("MID", "patientmedicaid"), 
    ("Address", "patientstreet"),
    ("Phone", "patientphone"),
    ("City", "patientcity"),
    ("State", "patientstate"),
    ("Zip Code","patientzip"),
    ("age", "patientage"), ("pat_phone", "patientphone"), ("physician", "physicianname")]

    # Fill Data
    all_fields = [(x[0], rform.get(x[1], "")) for x in all_fields_key]
    all_fields += [("DOB", datefix(rform.get('patientdob')))]

    # Fill Requested Equipment
    headers = ["hcpcs{}", "description{}", "modifier{}", "qty{}"]
    zheaders = ["hcpcs", "name", "modifier", "qty"]
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    firstpage_items = rl[:7]
    for index, record in enumerate(firstpage_items):
        fill_headers = map(lambda x: x.format(index + 1), headers)
        fill_kv = zip(fill_headers, [record.get(head, "") for head in zheaders])
        all_fields += fill_kv

    # Convert
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("forms/NV_DME_Prior_auth_form_data.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk forms/NV_DME_Prior_auth_form.pdf fill_form forms/NV_DME_Prior_auth_form_data.fdf output NV_DME_Prior_auth_form.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)
    zoho_upload("NV_DME_Prior_auth_form.pdf", rform.get("poid"))

# NV Mobility assessment and PA form
def nv_mobility_assessment_pa(rform):

    # Map fields
    all_fields_key = [
    ("7 Name last first", 'patientnamelf'), ("Address", "patientaddress"),
    ("9b Age", "patientage"), ("11 Phone", "patientphone"), 
    ("16 Name", "physicianname"), ("17 NPI", "physiciannpi"),
    ("18 Address include city state and zip", "physicianaddress"),
    ("19 Phone", "physicianphone"), ("20 Fax", "physicianfax")]

    # Fill Data
    all_fields = [(x[0], rform.get(x[1], "")) for x in all_fields_key]
    all_fields += [("9a Date of Birth", datefix(rform.get('patientdob')))]

    # Fill Requested Equipment
    headers = ["hcpcs{}", "description{}", "modifier{}", "qty{}"]
    zheaders = ["hcpcs", "name", "modifier","qty"]
    rl = createMapFromZohoString(rform.get('relatedrecord', ""))
    firstpage_items = rl[:19]
    for index, record in enumerate(firstpage_items):
        fill_headers = map(lambda x: x.format(index + 1), headers)
        fill_kv = zip(fill_headers, [record.get(head, "") for head in zheaders])
        all_fields += fill_kv

    # Convert
    fdf_data = fdfgen.forge_fdf("", all_fields, [], [], [])
    fdf_file = open("forms/NV_Mobility_assessment_and_PA_form_data.fdf", "w+")
    fdf_file.write(fdf_data)
    fdf_file.close()
    pdftk_cmd = "/app/vendor/pdftk/bin/pdftk forms/NV_Mobility_assessment_and_PA_form.pdf fill_form forms/NV_Mobility_assessment_and_PA_form_data.fdf output NV_Mobility_assessment_and_PA_form.pdf"
    os.chmod("/app/vendor/pdftk/bin/pdftk", 0o777)
    os.system(pdftk_cmd)
    zoho_upload("NV_Mobility_assessment_and_PA_form.pdf", rform.get("poid"))